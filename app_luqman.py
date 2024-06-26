from werkzeug.security import generate_password_hash, check_password_hash
from flask import Flask, render_template, request, redirect
from wtforms import StringField, SubmitField, IntegerField
from sqlalchemy import create_engine, text
from flask_wtf import FlaskForm
import pandas as pd
from docx import Document
from docx.shared import Cm
from dotenv import load_dotenv
import PIL
import datetime
import qrcode
from qrcode.image.styledpil import StyledPilImage
from qrcode.image.styles.moduledrawers.pil import RoundedModuleDrawer
from qrcode.image.styles.colormasks import RadialGradiantColorMask
import os

# Create engine

# load_dotenv()  # Loads the contents of the .env file into the environment as environment variables that we can access using os.getenv('VARIABLE_NAME')


# username = os.getenv("DB_USERNAME")
# password = os.getenv("DB_PASSWORD")
# host = os.getenv("DB_HOST")
# database = os.getenv("DB_NAME")
# path_to_excel_file = os.getenv("PRODUCT_DATA_PATH")
# quotation_dir = os.getenv("QUOTATION_DIR")

# PRODUCT_DATA_PATH=complete_product_list_spreadsheet.xlsx
# DB_USERNAME=root
# DB_PASSWORD=h1R1m!55
# DB_HOST=127.0.0.1
# DB_NAME=poof_schema
# QUOTATION_DIR =

username = "root"
password = "h1R1m!55"
host = "127.0.0.1"
database = "poof_schema"
path_to_excel_file = "complete_product_list_spreadsheet.xlsx"
quotation_dir = "quotations"

if not os.path.exists(quotation_dir):
    os.mkdir(quotation_dir)

codes = {}
current_quotation = []
current_client = {
    "Date": "",
    "Customer_Name": "",
    "Customer_Number": "",
    "Rep_Name": "",
    "Rep_Number": "",
}
customer_info = []

df = pd.read_excel(path_to_excel_file)
engine = create_engine(
    f"mysql+mysqlconnector://{username}:{password}@{host}/{database}"
)
df.to_sql(name="product_list", con=engine, if_exists="replace", index=False)
conn = engine.connect()

# Connect Flask
app = Flask(__name__)

# Create a document instance
quotation_doc = Document()

# Set Font of the document
style = quotation_doc.styles["Normal"]
style.font.name = "Arial"

# Add header to the document
header_section = quotation_doc.sections[0]
header = header_section.header
header_text = header.paragraphs[0]
header_text.text = "Multimedica ScO. All Rights Reserved"


@app.route("/", methods=["GET", "POST"])
def index():
    # TODO session tracking
    # If the user is not logged in, redirect to login page
    if request.method == "GET":
        return render_template("signIn.html")
    elif request.method == "POST":
        # Get username and password
        name = request.form.get("Username")
        Pass = request.form.get("Password")

        # Query Database for the data of the employee
        # TODO Protect against SQL injection
        authorityres = conn.execute(
            text(f'SELECT * FROM authorized_personnel WHERE Employee_Name = "{name}"')
        )
        conn.commit()
        authority = authorityres.all()

        # Ensure username and password were submitted
        if not name or not Pass:
            print("Name or Pass is empty")
            return render_template("invalid_credentials.html", name=name, Pass=Pass)

        # Ensure username exists and password is correct
        if len(authority) != 1 or not check_password_hash(authority[0][2], Pass):
            print("Wrong name or wrong Password")
            return render_template("invalid_credentials.html", authority=authorityres)

        # Redirect user to appropriate page depending on the authority level
        if authority[0][3] == "Administrator":
            return render_template("admin_options.html")
        elif authority[0][3] == "Data_Entry":
            return render_template("create_quotation.html")
        elif authority[0][3] == "Developer":
            return render_template("Developer_Options.html")
        else:
            return render_template("invalid_credentials.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    # TODO Make a change password page, hashes are great but now I don't know anyone's passwords and can't check them directly

    if request.method == "POST":
        # Get username and password and authority level

        name = request.form.get("username")
        password = request.form.get("up")
        Cpassword = request.form.get("up2")
        authority = request.form.get("authority")
        hash = generate_password_hash(password, method="pbkdf2", salt_length=16)
        # print(hash)
        if not name:
            return render_template("invalid_credentials.html")

        # Ensure password was submitted
        elif not password:
            return render_template("invalid_credentials.html")

        # Query database for username
        # TODO Protect against SQL injection
        result = conn.execute(
            text(f'SELECT * FROM authorized_personnel WHERE Employee_Name = "{name}"')
        )
        rows = result.all()

        # Ensure username exists and password is correct
        if len(rows) != 0:
            return render_template("invalid_credentials.html")

        if password != Cpassword:
            return render_template("invalid_credentials.html")

        # TODO Protect against SQL injection
        conn.execute(
            text(
                f"""INSERT INTO authorized_personnel(Employee_Name, Password, Authority_Level) 
                VALUES ("{name}", "{hash}", "{authority}")"""
            )
        )
        conn.commit()

        # Redirect user to home page
        return redirect("/")

    else:
        return render_template("register.html")


# Tasks page
@app.route("/task", methods=["GET", "POST"])
def task():
    if request.method == "GET":
        return render_template("admin_options.html")
    elif request.method == "POST":
        choice = request.form.get("task")
        # print(choice)
        rows = conn.execute(
            text(
                f"SELECT product_code, product_name FROM product_list ORDER BY product_code ASC"
            )
        )
        prods = rows.all()
        if choice == "create_quotation":
            # print(type(prods))
            return render_template("customer_info.html")
        elif choice == "Edit product prices":
            return render_template("edit_prices.html")
        elif choice == "Add new product":
            return render_template("add_product.html")
        elif choice == "Add Employee":
            return render_template("register.html")
        else:
            return render_template("invalid_choice.html")


@app.route("/customer_info", methods=["GET", "POST"])
def Customer_Info():
    if request.method == "POST":
        current_client["Date"] = request.form.get("quotation_date")
        current_client["Customer_Name"] = request.form.get("customer_name")
        current_client["Customer_Number"] = request.form.get("customer_number")
        current_client["Rep_Name"] = request.form.get("rep_name")
        current_client["Rep_Number"] = request.form.get("rep_number")
        rows = conn.execute(
            text(
                f"SELECT product_code, product_name FROM product_list ORDER BY product_code ASC"
            )
        )
        prods = rows.all()
        if (
            not current_client["Customer_Name"]
            or not current_client["Date"]
            or not current_client["Rep_Name"]
            or not current_client["Rep_Number"]
        ):
            return render_template("insufficient_data.html")
        return render_template(
            "create_quotation.html", products=prods, customer_info=current_client
        )
    pass


@app.route("/Edit_Quotation", methods=["GET", "POST"])
def Edit_Quotation():
    if request.method == "POST":
        rows = conn.execute(
            text(
                f"SELECT product_code, product_name FROM product_list ORDER BY product_code ASC"
            )
        )
        prods = rows.all()
        return render_template(
            "create_quotation.html",
            products=prods,
            customer_info=current_client,
            entries=current_quotation,
        )


# Create Quotation Page and handling Queries, autocomplete, and dynamic table row insertion
@app.route("/create_quotation", methods=["GET", "POST"])
def create_quotation():
    if request.method == "POST":
        # Get the product code and quantity from the submitted form
        code = request.form.get("product_code")
        quantity = request.form.get("quantity")
        quantity = float(quantity)
        rows = conn.execute(
            text(f"SELECT * FROM product_list WHERE product_code = '{code}'")
        )
        prod = rows.all()
        # redo the sql table and replace it with a list of lists in the frontend instead of a SQL table
        if len(prod) > 0:
            current_quotation.append(
                [
                    prod[0][4],
                    prod[0][0],
                    prod[0][1],
                    prod[0][3],
                    quantity,
                    prod[0][2],
                    prod[0][2] * quantity,
                ]
            )
        else:
            render_template("insufficient_data.html")
        """ print(prod[0][0])
        print(prod[0][1])
        print(prod[0][2])
        print(prod[0][3])
        print(prod[0][4]) """
        rows = conn.execute(
            text(
                f"SELECT product_code, product_name FROM product_list ORDER BY product_code ASC"
            )
        )
        prods = rows.all()
        return render_template(
            "create_quotation.html",
            entries=current_quotation,
            products=prods,
            customer_info=current_client,
        )


# TODO Add a route that allows the user to delete a product from the quotation. This will require a new form for each table row in the HTML


@app.route("/preview", methods=["GET", "POST"])
def preview():
    # TODO Add a preview page that shows the current quotation, an option to go back and edit, and an option to submit
    return render_template(
        "preview_quotation.html",
        customer_info=current_client,
        entries=current_quotation,
    )


# Convert the current quotation to a dataframe, submit it to the database, and export it as an excel file
@app.route("/export", methods=["GET", "POST"])
def submit():
    if request.method == "POST":
        client_columns = [
            "Date",
            "Customer_Name",
            "Customer_Number",
            "Rep_Name",
            "Rep_Number",
        ]
        product_columns = [
            "Image",
            "Product_Code",
            "Product_Name",
            "Description",
            "Quantity",
            "Price",
            "Total",
        ]
        client_list = [
            [
                current_client["Date"],
                current_client["Customer_Name"],
                current_client["Customer_Number"],
                current_client["Rep_Name"],
                current_client["Rep_Number"],
            ]
        ]
        client_data = pd.DataFrame(client_list, columns=client_columns)
        product_data = pd.DataFrame(current_quotation, columns=product_columns)

        print(product_data)
        # Export the dataframes to an Excel file, then save the file as a pdf
        # and save the pdf to the database alongside the name of the user and the date of submission
        # Add serializtion to the quotation files
        with pd.ExcelWriter("D:/Work/POOF/Quotation.xlsx") as writer:
            client_data.to_excel(writer, sheet_name="Client_Info", index=True)
            product_data.to_excel(writer, sheet_name="Product_Info", index=True)
        # Clear the current quotation and client info
        current_quotation.clear()
        current_client.clear()
        # Add Title
        Number = 1
        quotation_doc.add_heading(f"Quotation No. {Number}", 0)
        # Add Paragraphs
        p = quotation_doc.add_paragraph("Date: ")
        p.add_run(client_data["Date"][0])
        p = quotation_doc.add_paragraph("Customer Name: ")
        p.add_run(client_data["Customer_Name"][0])
        p = quotation_doc.add_paragraph("Customer Number: ")
        p.add_run(client_data["Customer_Number"][0])
        p = quotation_doc.add_paragraph("Rep Name: ")
        p.add_run(client_data["Rep_Name"][0])
        p = quotation_doc.add_paragraph("Rep Number: ")
        p.add_run(client_data["Rep_Number"][0])

        # Add QR Code of Quotation
        quotation_doc.add_picture(
            "D:/Work/POOF/qr_code.png", width=Cm(3.0), height=Cm(3.0)
        )

        # Add Table
        table = quotation_doc.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Image"
        hdr_cells[1].text = "Product Code"
        hdr_cells[2].text = "Product Name"
        hdr_cells[3].text = "Description"
        hdr_cells[4].text = "Quantity"
        hdr_cells[5].text = "Price"
        hdr_cells[6].text = "Total"
        for i in range(len(product_data)):
            row_cells = table.add_row().cells
            print(f"Product Code is: {product_data['Product_Code'][i]}")
            print(i)
            row_cells[0].text = str(product_data["Image"][i])
            row_cells[1].text = str(product_data["Product_Code"][i])
            row_cells[2].text = str(product_data["Product_Name"][i])
            row_cells[3].text = str(product_data["Description"][i])
            row_cells[4].text = str(product_data["Quantity"][i])
            row_cells[5].text = str(product_data["Price"][i])
            row_cells[6].text = str(product_data["Total"][i])
        quotation_doc.add_page_break()
        
        
        submit_quotation_to_db(session["user_id"], quotation_doc)
        
        
        # TODO Export the quotation to the database and save the pdf to the database
        # TODO Clear the document after exporting it to the database
        # Clear Dataframes
        client_data = pd.DataFrame()
        product_data = pd.DataFrame()
        print(client_data)
        print(product_data)
        return render_template("successful_submission.html"), {"Refresh": "10; url=/"}
    pass


# Create Page that allows admins to change the price of products
@app.route("/price", methods=["GET", "POST"])
def price():
    if request.method == "POST":
        choice = request.form.get("price")
        if choice == "single_product":
            return render_template("single_product.html")
        elif choice == "percentage":
            return render_template("percentage.html")
    return render_template("edit_prices.html")


# Change the price of a single product
@app.route("/single_price", methods=["GET", "POST"])
def single_price():
    if request.method == "POST":
        code = request.form.get("code")
        price = request.form.get("price")
        conn.execute(
            text(
                f'UPDATE product_list SET Price = "{price}" WHERE Product_Code = "{code}"'
            )
        )
        conn.commit()
        return render_template("edit_prices.html")
    return render_template("single_product.html")


# Change all product prices by a constant percentage either by increasing or decreasing
@app.route("/percentage", methods=["GET", "POST"])
def percentage():
    if request.method == "POST":
        percentage = request.form.get("percentage")
        Type = request.form.get("change_type")
        if not percentage:
            return render_template("invalid_choice.html")
        if Type == "Increase":
            new_percentage = 1 + (int(percentage) / 100)
        elif Type == "Decrease":
            new_percentage = 1 - (int(percentage) / 100)
        conn.execute(text(f"UPDATE product_list SET Price = Price * {new_percentage}"))
        conn.commit()
        return render_template("edit_prices.html")
    return render_template("percentage.html")


# Add a new product to the database
@app.route("/add_product", methods=["GET", "POST"])
def add_product():
    if request.method == "POST":
        code = request.form.get("code")
        name = request.form.get("name")
        price = request.form.get("price")
        description = request.form.get("description")
        image = request.form.get("image")  # TODO Convert image file to directory
        conn.execute(
            text(
                f'INSERT INTO product_list(Product_Code, Product_Name, Price, Description, Image_Directory) VALUES ("{code}", "{name}", "{price}", "{description}", "{image}")'
            )
        )
        conn.commit()
        return render_template("admin_options.html")
    return render_template("add_product.html")


################################################################################################
########################################### UTILS ##############################################
################################################################################################


# View the Quotation
@app.route("/view", methods=["GET"])  # expection a url/view?quotation_id=1
def view_quotation():
    quotation_id = request.args.get("quotation_id")
    
    file_path = conn.execute(
        text(
            f"SELECT quotation_file_path FROM exported_quotations WHERE quotation_id = {quotation_id}"
        )
    ).all()[0][0]
    
    
    ## TODO: Read the excel and put the values into the entries variable
    df = pd.read_excel(file_path)
    
    entires = [
        ["logo.png", 1, "Product 1", "Description 1", 500, 5, 2500],
        ["URL2", 2, "Product 2", "Description 2", 1000, 10, 10000],
        ["URL3", 3, "Product 3", "Description 3", 1500, 15, 22500],
        ["URL4", 4, "Product 4", "Description 4", 2000, 20, 40000],
        ["URL5", 5, "Product 5", "Description 5", 2500, 25, 62500],
        ["URL6", 6, "Product 6", "Description 6", 3000, 30, 90000],
        ["URL7", 7, "Product 7", "Description 7", 3500, 35, 122500],
    ]
    total = 0
    vat = 0
    for entry in entires:
        total += entry[6]
    vat = total * 0.14
    total += vat
    return render_template("view_quotation.html", entries=entires, total=total, vat=vat)


# Table that keeps record of exported files
# exported_quotations {
# 	(
# 		quotation_id: INT INCREMENTAL,
# 		submission_date: DATETIME,
# 		employee_id: INT,
# 		quotation_url: TEXT,
# 		quotation_file_path: TEXT,
# 	)
# }


"""
CREATE TABLE `poof_schema`.`exported_quotations` (
  `quotation_id` INT NOT NULL AUTO_INCREMENT,
  `submission_date` DATETIME NOT NULL,
  `employee_id` INT NOT NULL,
  `quotation_url` VARCHAR(255) NOT NULL,
  `quotation_file_path` VARCHAR(255) NOT NULL,
  PRIMARY KEY (`quotation_id`)
);
"""


# send data in format submit_quotation_to_db(1, "website.com/view/", "FILE_PATH")
def submit_quotation_to_db(employee_id: int, quotation: pd.DataFrame) -> bool:
    no_of_quotations = conn.execute(
        text("SELECT MAX(quotation_id) FROM exported_quotations")
    ).all()[0][0]

    quotation_id = no_of_quotations + 1
    submission_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    quotation_url = f"website.com/view?quotation_id={quotation_id}"  # TODO: Change this to the actual URL
    quotation_file_path = os.path.join(quotation_dir, f"{quotation_id}.xlsx")

    conn.execute(
        text(
            f"INSERT INTO exported_quotations(quotation_id, submission_date, employee_id, quotation_url, quotation_file_path) VALUES ({quotation_id}, '{submission_date}', {employee_id}, '{quotation_url}', '{quotation_file_path}')"
        )
    )
    conn.commit()

    with pd.ExcelWriter(quotation_file_path) as writer:
        quotation.to_excel(writer, sheet_name="Quotation", index=False)
    return True


def update_password(user_name: str, new_password) -> bool:
    # Make a SQL request to the DB to update the password hash for the user

    new_hash = generate_password_hash(new_password, method="pbkdf2", salt_length=16)

    query = f"UPDATE authorized_personnel SET Password = '{new_hash}' WHERE Employee_Name = '{user_name}'"
    response = conn.execute(text(query))

    return (
        response.rowcount == 1
    )  # If the row was updated, return True, else return False


# Return the image of the QR code to the calling function
def convert_url_to_qr_code(url: str, rounded_corners=True, logo_path=None) -> PIL.Image:
    qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_L)
    qr.add_data(url)

    if rounded_corners and logo_path:
        img = qr.make_image(
            image_factory=StyledPilImage,
            module_drawer=RoundedModuleDrawer(),
            embeded_image_path=logo_path,
        )
    elif rounded_corners:
        img = qr.make_image(
            image_factory=StyledPilImage, module_drawer=RoundedModuleDrawer()
        )
    elif logo_path:
        img = qr.make_image(image_factory=StyledPilImage, embeded_image_path=logo_path)
    else:
        img = qr.make_image()

    return img


""" 
# Unused Route, replaced, delete later
@app.route("/query", methods=["GET", "POST"])
def query():
    if request.method == "GET":
        return render_template("create_quotation.html")
    elif request.method == "POST":
        code = request.form.get("product_code")
        name = request.form.get("product_name")
        print(name)
        print(code)
        quantity = request.form.get("quantity")
        print(quantity)
        if (not code and not name) or not quantity:
            return render_template("insufficient_data.html")
        product_details = get_product(code, name, quantity)

        return render_template("create_quotation.html", product_details=product_details, quantity=quantity)
 """


"""
# Unused Route, replaced, delete later
def get_date():
    print("Please enter the date: ")
    date = input()
    return date
 """

""" 
# Unused Route, replaced, delete later
def get_product(code, name, quantity):
    pname = name
    pcode = code
    
    if not pcode:
        pcode = c.execute(text(f"SELECT Product_Code FROM product_list WHERE Product_Name = '{pname}'"))
        pcode = pcode.all()
    if not pname:
        pname = c.execute(text(f"SELECT Product_Name FROM product_list WHERE Product_Code = '{pcode}'"))
        pname = pname.all()
    pprice = c.execute(text(f"SELECT Price FROM product_list WHERE product_code = '{pcode}'"))
    pprice = pprice.all()

    pdescription = c.execute(text(f"SELECT Description FROM product_list WHERE product_code = '{pcode}'"))
    pdescription = pdescription.all()
    pimg_dir = c.execute(text(f"SELECT Image_Directory FROM product_list WHERE product_code = '{pcode}'"))
    pimg_dir = pimg_dir.all()
    product = [pimg_dir, pcode, pname, pdescription, quantity, pprice]
    return product
 """
""" 
# Unused Route, replaced, delete later
def get_quantity():
    print("Please enter the quantity: ")
    quantity = input()
    return quantity
 """
""" 
# Unused Route, replaced, delete later
def get_client_name():
    print("Please enter the client name: ")
    client = input()
    return client """
